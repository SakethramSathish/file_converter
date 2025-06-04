import streamlit as st
from docx import Document
from docx2pdf import convert as docx2pdf_convert
from pdf2docx import Converter
import os
import tempfile

def convert():
    option = st.selectbox("Select Conversion Type", [
        "DOCX to PDF",
        "PDF to DOCX",
        "TXT to DOCX",
        "DOCX to TXT",
        "ODT to DOCX",
        "ODT to PDF"
    ])

    uploaded_file = st.file_uploader("Upload your file", type=["docx", "pdf", "txt", "odt"])

    if uploaded_file and st.button("Convert"):
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = os.path.join(tmpdir, uploaded_file.name)
            with open(input_path, "wb") as f:
                f.write(uploaded_file.read())

            if option == "DOCX to PDF":
                output_path = os.path.join(tmpdir, "output.pdf")
                docx2pdf_convert(input_path, output_path)

            elif option == "PDF to DOCX":
                output_path = os.path.join(tmpdir, "output.docx")
                cv = Converter(input_path)
                cv.convert(output_path, start=0, end=None)
                cv.close()

            elif option == "TXT to DOCX":
                output_path = os.path.join(tmpdir, "output.docx")
                doc = Document()
                with open(input_path, "r", encoding="utf-8") as f:
                    for line in f:
                        doc.add_paragraph(line.read())
                doc.save(output_path)

            elif option == "DOCX to TXT":
                output_path = os.path.join(tmpdir, "output.txt")
                doc = Document(input_path)
                with open(output_path, "w", encoding="utf-8") as f:
                    for para in doc.paragraphs:
                        f.write(para.text + "\n")

            elif option.startswith("ODT"):
                output_path = os.path.join(tmpdir, "output." + ("pdf" if "PDF" in option else "docx"))
                os.system(f'libreoffice --headless --convert-to {"pdf" if "PDF" in option else "docx"} "{input_path}" --outdir "{tmpdir}"')

            with open(output_path, "rb") as f:
                st.download_button("Download Converted File", f, file_name=os.path.basename(output_path))
            