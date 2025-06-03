import os
from docx2pdf import convert as docx2pdf_convert
from pdf2docx import Converter
from docx import Document

def docx_to_pdf(input_path, output_path=None):
    """
    Convert a DOCX file to PDF format.
    
    """
    if output_path is None:
        output_path = os.path.splitext(input_path)[0] + '.pdf'
    
    try:
        docx2pdf_convert(input_path, output_path)
        return output_path
    except Exception as e:
        raise RuntimeError(f"Failed to convert DOCX to PDF: {e}")
    
def pdf_to_docx(input_path, output_path):
    """
    Convert a PDF file to DOCX format.
    
    """
    if output_path is None:
        output_path = os.path.splitext(input_path)[0] + '.docx'
    
    try:
        cv = Converter(input_path)
        cv.convert(output_path, start=0, end=None)
        cv.close()
        return output_path
    except Exception as e:
        raise RuntimeError(f"Failed to convert PDF to DOCX: {e}")
    
def txt_to_docx(input_path, output_path):
    """
    Convert a TXT file to DOCX format.
    
    """
    if output_path is None:
        output_path = os.path.splitext(input_path)[0] + '.docx'
    
    try:
        doc = Document()
        with open(input_path, 'r', encoding='utf-8') as file:
            for line in file:
                doc.add_paragraph(line.strip())
        doc.save(output_path)
        return output_path
    except Exception as e:
        raise RuntimeError(f"Failed to convert TXT to DOCX: {e}")
    
def docx_to_txt(input_path, output_path=None):
    """
    Convert a DOCX file to TXT format.
    
    """
    if output_path is None:
        output_path = os.path.splitext(input_path)[0] + '.txt'
    
    try:
        doc = Document(input_path)
        with open(output_path, 'w', encoding='utf-8') as file:
            for para in doc.paragraphs:
                file.write(para.text + '\n')
        return output_path
    except Exception as e:
        raise RuntimeError(f"Failed to convert DOCX to TXT: {e}")